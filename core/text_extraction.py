"""
text_extraction.py
-------------------
Extracts raw text from resumes / job descriptions supplied as .pdf, .docx,
or .txt files. Designed to accept either a filesystem path or an in-memory
file-like object (so it works both from the CLI script and from Streamlit's
file uploader, which hands back BytesIO-like objects).
"""

import io
import os

import pdfplumber
import docx  # python-docx


class UnsupportedFileTypeError(Exception):
    pass


def _get_extension(file_obj_or_path) -> str:
    """Works for both a path string and an uploaded file object with .name"""
    if isinstance(file_obj_or_path, str):
        name = file_obj_or_path
    else:
        name = getattr(file_obj_or_path, "name", "")
    return os.path.splitext(name)[1].lower()


def extract_text(file_obj_or_path) -> str:
    """
    Extract raw text from a resume/JD file.

    Parameters
    ----------
    file_obj_or_path : str | file-like object
        A filesystem path OR an uploaded file object (e.g. from
        Streamlit's `st.file_uploader`, which behaves like BytesIO but
        also exposes a `.name` attribute).

    Returns
    -------
    str : extracted plain text (empty string if nothing could be read)
    """
    ext = _get_extension(file_obj_or_path)

    # Defensive: rewind file-like objects (e.g. Streamlit's UploadedFile) to
    # the start in case they were already read earlier in the same run —
    # otherwise a second read silently returns empty text.
    if not isinstance(file_obj_or_path, str) and hasattr(file_obj_or_path, "seek"):
        file_obj_or_path.seek(0)

    if ext == ".pdf":
        return _extract_pdf(file_obj_or_path)
    elif ext == ".docx":
        return _extract_docx(file_obj_or_path)
    elif ext == ".txt":
        return _extract_txt(file_obj_or_path)
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext}'. Supported types: .pdf, .docx, .txt"
        )


def _extract_pdf(file_obj_or_path) -> str:
    text_chunks = []
    with pdfplumber.open(file_obj_or_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def _extract_docx(file_obj_or_path) -> str:
    document = docx.Document(file_obj_or_path)
    paragraphs = [p.text for p in document.paragraphs]
    # Also capture text inside tables (common in resume templates)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _extract_txt(file_obj_or_path) -> str:
    if isinstance(file_obj_or_path, str):
        with open(file_obj_or_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raw = file_obj_or_path.read()
        if isinstance(raw, bytes):
            return raw.decode("utf-8", errors="ignore")
        return raw


def get_candidate_name(file_obj_or_path, fallback_index: int = 0) -> str:
    """
    Best-effort candidate name derived from the filename
    (e.g. 'john_doe_resume.pdf' -> 'John Doe Resume').
    Real name extraction from resume content is intentionally left out --
    it's unreliable with regex/NLP alone and out of scope for this project;
    the filename is a transparent, editable stand-in used consistently
    throughout the ranked output.
    """
    if isinstance(file_obj_or_path, str):
        name = os.path.basename(file_obj_or_path)
    else:
        name = getattr(file_obj_or_path, "name", f"Candidate_{fallback_index}")

    name = os.path.splitext(name)[0]
    name = name.replace("_", " ").replace("-", " ").strip()
    return name.title() if name else f"Candidate {fallback_index}"
